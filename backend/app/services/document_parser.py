"""Document text extraction for various file types."""

import io


async def extract_text(content: bytes, filename: str, mime_type: str) -> str:
    """Extract text from uploaded document. Supports PDF, DOCX, TXT, Markdown, CSV."""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    # Plain text files
    if ext in ("txt", "md", "csv", "tsv", "log", "json", "xml", "html", "sql", "ddl") or mime_type.startswith("text/"):
        try:
            return content.decode("utf-8")
        except UnicodeDecodeError:
            return content.decode("euc-kr", errors="replace")

    # PDF
    if ext == "pdf" or mime_type == "application/pdf":
        return _extract_pdf(content)

    # DOCX
    if ext == "docx" or mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return _extract_docx(content)

    # YAML files - return as-is
    if ext in ("yaml", "yml"):
        try:
            return content.decode("utf-8")
        except UnicodeDecodeError:
            return content.decode("euc-kr", errors="replace")

    # Fallback: try as text
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        raise ValueError(f"Unsupported file type: {ext} ({mime_type}). Supported: PDF, DOCX, TXT, MD, CSV, SQL, DDL, YAML")


def _extract_pdf(content: bytes) -> str:
    """Extract text from PDF bytes."""
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        raise ValueError("PDF support requires PyPDF2. Install with: pip install PyPDF2")

    reader = PdfReader(io.BytesIO(content))
    texts = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            texts.append(text)

    if not texts:
        raise ValueError("PDF에서 텍스트를 추출할 수 없습니다. 스캔된 이미지 PDF일 수 있습니다.")

    return "\n\n".join(texts)


def _extract_docx(content: bytes) -> str:
    """Extract text from DOCX bytes."""
    try:
        from docx import Document
    except ImportError:
        raise ValueError("DOCX support requires python-docx. Install with: pip install python-docx")

    doc = Document(io.BytesIO(content))
    texts = []

    for para in doc.paragraphs:
        if para.text.strip():
            texts.append(para.text)

    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                texts.append(row_text)

    if not texts:
        raise ValueError("DOCX에서 텍스트를 추출할 수 없습니다.")

    return "\n".join(texts)
